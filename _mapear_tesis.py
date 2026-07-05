"""Mapea la estructura del docx para encontrar posiciones clave."""
from docx import Document
import os, io, sys

sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r"C:\Users\HP\OneDrive\Escritorio\edwin\TESIS_20212444.docx")

# Palabras clave a detectar
patrones = {
    'RESUMEN': [],
    'Introducción': [],
    'Introducción:': [],
    'Capítulo 1': [],
    'Capítulo 2': [],
    'Capítulo 3': [],
    'Capítulo 4': [],
    '1. Estudio': [],
    '1.5': [],
    '1.5.1': [],
    '1.5.2': [],
    '2.1': [],
    '2. Fundamentos': [],
    '3. Diseño': [],
    '3.1': [],
    '4.': [],
    '4.1': [],
    '4.2': [],
    '4.3': [],
    '4.4': [],
    '4.5': [],
    '4.6': [],
    'Conclusiones': [],
    'Bibliografía': [],
    'Referencias': [],
    'REFERENCIAS': [],
    'Anexo': [],
    'ANEXOS': [],
    '[1]': [],
    '[10]': [],
    '[20]': [],
}

for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if not txt:
        continue
    for pat in patrones:
        if txt.startswith(pat):
            patrones[pat].append((i, txt[:120]))
            break

for pat, hits in patrones.items():
    print(f"\n=== {pat} ===")
    for idx, txt in hits[:5]:
        print(f"  [{idx:04d}] {txt}")
    if len(hits) > 5:
        print(f"  ... ({len(hits)-5} mas)")

print(f"\n\nTotal parrafos: {len(doc.paragraphs)}")
