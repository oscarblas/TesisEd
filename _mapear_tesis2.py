"""Mapa mas fino de la tesis."""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r"C:\Users\HP\OneDrive\Escritorio\edwin\TESIS_20212444.docx")

# Ver rangos criticos
print("=== PARRAFOS 150-250 (contenido de intro y Cap 1) ===")
for i, p in enumerate(doc.paragraphs[150:250], start=150):
    t = p.text.strip()
    if t:
        print(f"[{i:04d}] {t[:130]}")

print("\n=== PARRAFOS 550-700 (Cap 4 esperado) ===")
for i, p in enumerate(doc.paragraphs[550:700], start=550):
    t = p.text.strip()
    if t:
        print(f"[{i:04d}] {t[:130]}")

print("\n=== PARRAFOS 700-812 (final: Bibliografia + Anexos) ===")
for i, p in enumerate(doc.paragraphs[700:812], start=700):
    t = p.text.strip()
    if t:
        print(f"[{i:04d}] {t[:130]}")
