"""
Genera TESIS_20212444_secciones_nuevas.docx con las secciones
que Edwin debe insertar/reemplazar en su tesis principal.

Contenido:
- Introducción actualizada
- Objetivos (general + especificos) actualizados
- Capitulo 4 completo redactado
- Conclusiones (capitulo)
- Bibliografia formato APA (segun instructivo PUCP)
- Anexos A-E con codigos MATLAB

Formato: Times New Roman 12, doble espacio, margenes 2.54 cm,
indentado de 1 cm al inicio de cada parrafo.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"C:\Users\HP\OneDrive\Escritorio\edwin"

# ============== CONFIGURACION DE FORMATO PUCP ==============
FONT       = "Times New Roman"
FONT_SIZE  = Pt(12)
MARGIN     = Cm(2.54)
INDENT     = Cm(1.0)
LINE_SPACE = 2.0

COURIER      = "Courier New"
CODE_SIZE    = Pt(9)

doc = Document()

# --- Configurar margenes ---
for section in doc.sections:
    section.top_margin    = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin   = MARGIN
    section.right_margin  = MARGIN

# --- Configurar estilo Normal ---
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = FONT_SIZE
pf = style.paragraph_format
pf.line_spacing = LINE_SPACE
pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
pf.space_after = Pt(0)
pf.space_before = Pt(0)

# ============== UTILIDADES ==============

def add_paragraph(text, bold=False, italic=False, align=None, indent=True,
                  space_after=0, size=None, center=False, uppercase=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    pf = p.paragraph_format
    if indent and not center:
        pf.first_line_indent = INDENT
    pf.space_after = Pt(space_after)

    r = p.add_run(text.upper() if uppercase else text)
    r.font.name = FONT
    r.font.size = size if size else FONT_SIZE
    r.bold = bold
    r.italic = italic
    return p

def add_capitulo_titulo(numero, titulo):
    """Nivel 1: Centrado, negrita, mayusculas y minusculas."""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(f"Capítulo {numero}")
    r.font.name = FONT; r.font.size = Pt(14); r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(titulo)
    r.font.name = FONT; r.font.size = Pt(14); r.bold = True

def add_seccion_nivel2(numero, titulo):
    """Nivel 2: Justificado a la izquierda, negrita."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"{numero} {titulo}")
    r.font.name = FONT; r.font.size = Pt(12); r.bold = True

def add_seccion_nivel3(numero, titulo):
    """Nivel 3: Indentado, negrita, primera letra mayuscula."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{numero} {titulo}.")
    r.font.name = FONT; r.font.size = Pt(12); r.bold = True

def add_formula_box(formula_desc, latex_code):
    """Bloque para insertar una formula manualmente."""
    # Descriptor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"[Insertar en Word con editor de ecuaciones — modo LaTeX]")
    r.font.name = FONT; r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RGBColor(0x80,0x80,0x80)

    # Codigo LaTeX
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r = p.add_run(latex_code)
    r.font.name = COURIER; r.font.size = Pt(11)

def add_marcador_imagen(texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"[{texto}]")
    r.font.name = FONT; r.font.size = Pt(11); r.italic = True; r.font.color.rgb = RGBColor(0xC0,0x39,0x2B)

def add_codigo_matlab(nombre_archivo, ruta):
    """Inserta el codigo con formato monoespaciado, espacio simple."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"Archivo: {nombre_archivo}")
    r.font.name = FONT; r.font.size = Pt(11); r.bold = True; r.italic = True

    if not os.path.exists(ruta):
        p = doc.add_paragraph()
        r = p.add_run(f"[Archivo no encontrado: {ruta}]")
        r.font.name = FONT; r.italic = True
        return

    with open(ruta, 'r', encoding='utf-8', errors='replace') as f:
        contenido = f.read()

    for linea in contenido.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(linea if linea else " ")
        r.font.name = COURIER; r.font.size = CODE_SIZE

def add_bibliografia_entry(texto):
    """Sangria francesa: primera linea sin indentar, siguientes 1.27cm."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(1.27)
    pf.first_line_indent = Cm(-1.27)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.name = FONT; r.font.size = Pt(12)

# ============================================================
# ADVERTENCIA INICIAL
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SECCIONES NUEVAS PARA LA TESIS")
r.font.name = FONT; r.font.size = Pt(16); r.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Reemplazar / insertar en TESIS_20212444.docx según indicaciones")
r.font.name = FONT; r.font.size = Pt(12); r.italic = True

doc.add_paragraph()

add_paragraph(
    "Este archivo contiene todas las secciones nuevas o modificadas que deben integrarse "
    "en la tesis principal. Todas las secciones ya están redactadas con el formato PUCP "
    "(Times New Roman 12, doble espacio, márgenes 2.54 cm, párrafos justificados con "
    "indentado de 1 cm). Las ecuaciones se presentan como código LaTeX listo para insertar "
    "con el editor de ecuaciones de Word (Alt = , modo LaTeX). Las figuras se marcan con "
    "[IMAGEN X.Y] para adjuntarlas desde los scripts MATLAB del repositorio.",
    indent=False, space_after=12
)

# ============================================================
# SECCION 1: INTRODUCCION ACTUALIZADA
# ============================================================
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Introducción")
r.font.name = FONT; r.font.size = Pt(14); r.bold = True

add_paragraph(
    "El control automático de nivel en sistemas de tanques acoplados constituye uno "
    "de los desafíos clásicos y a la vez más recurrentes en la ingeniería de procesos "
    "industriales. Se encuentra presente en sectores tan diversos como el petroquímico, "
    "el farmacéutico, el minero, el alimentario y el energético, donde una regulación "
    "precisa del nivel de los tanques resulta indispensable para garantizar la calidad "
    "del producto final, la seguridad de la operación y la eficiencia del proceso. "
    "La complejidad de esta tarea se incrementa notablemente cuando los tanques se "
    "encuentran acoplados entre sí, dado que las variables del sistema dejan de ser "
    "independientes y las estrategias de control convencionales pierden efectividad.",
)

add_paragraph(
    "El sistema de cuatro tanques acoplados, propuesto por Karl Henrik Johansson a "
    "finales de los años noventa, se ha consolidado como una plataforma de referencia "
    "para el estudio de procesos multivariables. Su principal virtud radica en que, "
    "mediante la variación de un único parámetro físico, es posible modificar la posición "
    "de los ceros del sistema, lo que permite estudiar en un mismo montaje comportamientos "
    "de fase mínima y fase no mínima. Esta propiedad lo convierte en un banco de pruebas "
    "ideal para evaluar estrategias de control clásico y avanzado en un mismo entorno "
    "experimental.",
)

add_paragraph(
    "Frente a las limitaciones de los controladores clásicos ante procesos multivariables "
    "con acoplamiento cruzado, restricciones físicas y operación fuera del punto nominal, "
    "el Control Predictivo Basado en Modelo (MPC, por sus siglas en inglés) se ha erigido "
    "como una alternativa robusta y versátil. Dentro de esta familia, el Control Predictivo "
    "Generalizado (GPC) destaca por su formulación matemática elegante, su capacidad para "
    "gestionar de manera natural los sistemas multivariables mediante el modelo CARIMA y "
    "las ecuaciones diofánticas, y por su compatibilidad con la incorporación explícita de "
    "restricciones a través de una formulación de programación cuadrática.",
)

add_paragraph(
    "La presente investigación se enfoca en el diseño, la implementación en un entorno de "
    "simulación y el análisis comparativo de un controlador GPC multivariable aplicado al "
    "sistema hidráulico de cuatro tanques acoplados. El desempeño del controlador diseñado "
    "se contrasta con el de un esquema de control PI descentralizado con desacoplador "
    "estático, considerado la mejor configuración clásica practicable para un sistema TITO "
    "acoplado en la industria. La comparación se sustenta en un escenario integrado de "
    "simulación que combina, en una sola secuencia operativa, los principales desafíos "
    "del control multivariable: arranque desde tanques vacíos, cambios secuenciales y "
    "simultáneos de referencia, inyección de ruido gaussiano en los sensores y operación "
    "significativamente alejada del punto de linealización.",
)

add_paragraph(
    "El documento se organiza en cuatro capítulos. El Capítulo 1 presenta el estudio del "
    "control de nivel en sistemas de tanques acoplados, incluyendo la descripción del "
    "sistema, el estado del arte, la justificación del estudio y los objetivos de la "
    "investigación. El Capítulo 2 desarrolla los fundamentos teóricos del Control Predictivo "
    "Generalizado y el modelado matemático del sistema de cuatro tanques. El Capítulo 3 "
    "aborda el diseño del controlador GPC MIMO, incluyendo la selección de sus parámetros, "
    "la construcción de sus matrices características, el tratamiento de restricciones y "
    "el análisis comparativo de cuatro métodos de sintonización. El Capítulo 4 desarrolla "
    "el análisis comparativo del GPC diseñado frente al controlador PI con desacoplador. "
    "Finalmente, se presentan las conclusiones generales del trabajo, las recomendaciones "
    "para líneas futuras de investigación, la bibliografía consultada y los anexos que "
    "contienen los códigos MATLAB desarrollados.",
    space_after=12
)

# ============================================================
# SECCION 2: OBJETIVOS ACTUALIZADOS
# ============================================================
doc.add_page_break()
add_seccion_nivel2("1.5", "Objetivos del trabajo de investigación")

add_seccion_nivel3("1.5.1", "Objetivo general")

add_paragraph(
    "Diseñar e implementar en un entorno de simulación un controlador predictivo "
    "generalizado (GPC) multivariable para el sistema hidráulico de cuatro tanques "
    "acoplados, orientado a la regulación coordinada de los niveles de los tanques "
    "inferiores en presencia de acoplamiento cruzado, perturbaciones externas y "
    "restricciones físicas sobre los actuadores, garantizando un desempeño superior "
    "al de las estrategias de control clásicas empleadas habitualmente en la industria.",
)

add_seccion_nivel3("1.5.2", "Objetivos específicos")

for i, obj in enumerate([
    "Caracterizar la dinámica multivariable del sistema de cuatro tanques acoplados, "
    "identificando sus propiedades estructurales, sus modos de operación de fase mínima "
    "y fase no mínima, y los caminos directos y cruzados entre entradas y salidas.",

    "Obtener el modelo linealizado del sistema en espacio de estados y su "
    "correspondiente matriz de funciones de transferencia MIMO en torno a un punto de "
    "operación nominal, y validar la fidelidad del modelo lineal respecto al modelo "
    "no lineal mediante la métrica de ajuste porcentual (FIT%).",

    "Diseñar el algoritmo del controlador Predictivo Generalizado para el sistema TITO, "
    "lo cual incluye la construcción de la matriz dinámica, la definición de los "
    "horizontes de predicción y control, la ponderación de la función de costo y el "
    "tratamiento de las restricciones físicas mediante una formulación de programación "
    "cuadrática.",

    "Implementar el controlador GPC diseñado en el entorno MATLAB/Simulink y evaluar "
    "su desempeño mediante los criterios integrales del error, el tiempo de "
    "establecimiento, el sobrepico y el esfuerzo de control.",

    "Realizar un análisis comparativo entre el desempeño del controlador GPC diseñado "
    "y el de un controlador PI descentralizado con desacoplador estático sintonizado "
    "por Internal Model Control, identificando las condiciones bajo las cuales el "
    "GPC justifica su adopción frente a la alternativa clásica."
]):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i+1}. {obj}")
    r.font.name = FONT; r.font.size = Pt(12)

# ============================================================
# SECCION 3: CAPITULO 4 COMPLETO
# ============================================================
add_capitulo_titulo(4,
    "Análisis Comparativo del Controlador GPC frente al Control PI con Desacoplador "
    "en el Sistema de Cuatro Tanques Acoplados")

# 4.1 Introduccion
add_seccion_nivel2("4.1", "Introducción")

add_paragraph(
    "En el presente capítulo se desarrolla un análisis comparativo entre un controlador "
    "clásico con el controlador GPC desarrollado en el capítulo anterior, implementado "
    "sobre la misma planta no lineal de cuatro tanques acoplados. El diseño para la "
    "comparativa del controlador de contraste busca evidenciar las diferentes respuestas "
    "del sistema con control clásico y con algoritmos de control avanzado GPC en nuestro "
    "sistema industrial.",
)

add_paragraph(
    "Para la comparativa, se ha optado por diseñar un controlador PI en vez de un PID. "
    "Esto debido a que en el sistema de cuatro tanques, la dinámica vista por cada lazo "
    "es dominantemente de primer orden con constante de tiempo del orden de las decenas "
    "de segundos. En estas condiciones, el aporte del término derivativo es marginal, "
    "mientras que su efecto sobre la amplificación del ruido de los transmisores de "
    "presión es considerable. Por esta razón, la práctica industrial recomienda omitir "
    "el término derivativo en aplicaciones de control de nivel y trabajar con la "
    "formulación PI (T_d = 0) [24].",
)

add_paragraph(
    "Existen estrategias aplicables a un controlador clásico PI para enfrentarnos a un "
    "sistema multivariable, como por ejemplo el PI descentralizado puro (sin desacoplador), "
    "que deja sin compensar la interacción cruzada entre lazos, lo que en un sistema TITO "
    "acoplado es significativo y es justo lo que se desea disminuir, puesto que el cambio "
    "de una referencia provoca una perturbación inmediata en la salida opuesta. La "
    "estrategia industrial estándar consiste en incorporar un desacoplador estático entre "
    "las salidas de los PI y las entradas a la planta para cancelar la interacción del "
    "acople en estado estacionario, mejorando notablemente el comportamiento global del "
    "lazo. Sin este componente la comparación con el GPC sería ventajosamente sesgada "
    "hacia el predictivo; con desacoplador, en cambio, se garantiza que el contraste se "
    "realice contra la mejor configuración clásica practicable.",
)

add_paragraph(
    "El capítulo se organiza en seis secciones. La sección 4.2 desarrolla el diseño "
    "completo del PI con desacoplador. La sección 4.3 describe el escenario integrado "
    "de simulación. La sección 4.4 reporta el análisis cuantitativo de los resultados. "
    "La sección 4.5 discute los hallazgos y valida la hipótesis del trabajo. La sección "
    "4.6 cierra con las conclusiones del capítulo.",
)

# 4.2 Diseño del controlador PI con desacoplador estático
add_seccion_nivel2("4.2", "Diseño del controlador PI con desacoplador estático")

# 4.2.1
add_seccion_nivel3("4.2.1", "Emparejamiento entrada–salida")

add_paragraph(
    "En el sistema de cuatro tanques acoplados, cada bomba afecta directamente al tanque "
    "inferior de su rama y, de manera indirecta, al tanque inferior opuesto a través del "
    "acoplamiento cruzado de los tanques superiores. Al examinar la matriz B_c del modelo "
    "linealizado se identifican los caminos directos, que son los rápidos: u₁ → h₄ con "
    "ganancia γ₁·k₁/A₄ y u₂ → h₃ con ganancia γ₂·k₂/A₃. En consecuencia, el emparejamiento "
    "natural adoptado es que PI_1 controle h₄ a través de u₁, y que PI_2 controle h₃ a "
    "través de u₂. Esta selección coincide con la recomendada por Johansson [8] para "
    "configuraciones de fase mínima (γ₁+γ₂ > 1), que es justamente el caso de la planta "
    "del Laboratorio de Control Avanzado de la PUCP.",
)

# 4.2.2
add_seccion_nivel3("4.2.2", "Sintonización por Internal Model Control (IMC)")

add_paragraph(
    "Cada PI se diseña de forma independiente, considerando que la otra entrada permanece "
    "constante en su valor estacionario. Bajo este supuesto, el subproceso visto por cada "
    "lazo se aproxima a un sistema de primer orden con ganancia K y constante de tiempo τ:",
)

add_formula_box("Subproceso primer orden", r"G_{loop}(s) = \frac{K}{\tau\,s + 1}")

add_paragraph(
    "Se ha optado por el método de Internal Model Control debido a que sus reglas son "
    "analíticas y cerradas, lo cual elimina la subjetividad de los métodos heurísticos y "
    "permite una sintonización reproducible para cualquier planta del mismo tipo. "
    "Adicionalmente, la sintonización queda parametrizada por un único valor λ_imc que "
    "posee un significado físico directo como constante de tiempo deseada en lazo cerrado, "
    "característica que facilita el ajuste por parte del personal técnico. Los parámetros "
    "del PI resultan:",
)

add_formula_box("Parametros PI", r"K_p = \frac{\tau}{K \cdot \lambda_{imc}},\quad T_i = \tau")

add_paragraph(
    "Para el valor de λ_imc se adopta la regla λ_imc = τ/3, recomendada por la literatura "
    "industrial como ajuste por defecto en procesos de nivel y temperatura [17]. Valores "
    "menores que τ/3 aceleran el seguimiento, mientras que erosionan los márgenes de "
    "estabilidad; valores mayores incrementan la robustez, en cambio elevan considerablemente "
    "el tiempo de establecimiento. La regla adoptada constituye un compromiso entre ambos "
    "extremos.",
)

add_marcador_imagen("IMAGEN 4.1 — Respuesta en lazo cerrado del PI para tres valores de λ_imc (τ/2, τ/3, τ). Generada con analisis_lambda_imc.m")

add_marcador_imagen("TABLA 4.1 — Parámetros de los PI descentralizados sintonizados por IMC. Generada con controlador_PID.m")

# 4.2.3
add_seccion_nivel3("4.2.3", "Desacoplador estático")

add_paragraph(
    "Para mitigar el efecto del acoplamiento cruzado, se incorpora un desacoplador estático "
    "entre las salidas de los PI y las entradas a la planta. Se adopta la formulación "
    "simplificada de Skogestad [19], que mantiene la diagonal unitaria y emplea las "
    "ganancias DC cruzadas para cancelar la interacción:",
)

add_formula_box("Matriz D",
    r"\mathbf{D} = \begin{bmatrix} 1 & -k_{12} \\ -k_{21} & 1 \end{bmatrix}")

add_paragraph(
    "Esta formulación se prefiere a un desacoplador dinámico debido a que requiere "
    "únicamente la matriz de ganancias DC del modelo, mucho más sencilla de obtener y "
    "notablemente más robusta a la incertidumbre paramétrica que un modelo dinámico "
    "completo. Adicionalmente, la cancelación estacionaria suele ser suficiente cuando "
    "los lazos individuales son moderadamente rápidos respecto al acople cruzado, "
    "condición que se cumple en el sistema piloto bajo estudio.",
)

add_paragraph(
    "Los coeficientes k₁₂ y k₂₁ se calculan a partir de la matriz de ganancias DC del "
    "sistema linealizado:",
)

add_formula_box("G_dc", r"\mathbf{G}_{dc} = -\mathbf{C}_c \mathbf{A}_c^{-1} \mathbf{B}_c")

add_formula_box("Coeficientes desacoplador",
    r"k_{12} = \frac{G_{dc}(h_4, u_2)}{G_{dc}(h_4, u_1)},\quad "
    r"k_{21} = \frac{G_{dc}(h_3, u_1)}{G_{dc}(h_3, u_2)}")

add_paragraph(
    "La aplicación del desacoplador convierte las salidas v = [v_1, v_2]ᵀ de los PI en "
    "las señales u = [u_1, u_2]ᵀ que efectivamente alimentan las bombas:",
)

add_formula_box("Aplicacion D", r"\mathbf{u} = \mathbf{D}\,\mathbf{v}")

add_paragraph(
    "Conviene anticipar una limitación que se verá reflejada en los resultados: el "
    "desacoplador estático cancela el acople únicamente en estado estacionario y "
    "alrededor del punto de operación nominal. Cuando el sistema opera en regiones "
    "alejadas, las ganancias reales de la planta cambian debido a la dependencia con "
    "√h, mientras que el desacoplador conserva sus coeficientes fijados a partir del "
    "punto nominal y, en consecuencia, pierde efectividad.",
)

add_marcador_imagen("TABLA 4.2 — Coeficientes del desacoplador. Generada con controlador_PID.m")

# 4.2.4
add_seccion_nivel3("4.2.4", "Algoritmo discreto con anti-windup")

add_paragraph(
    "La implementación digital del controlador se realiza en forma incremental, también "
    "llamada forma de velocidad:",
)

add_formula_box("Forma incremental",
    r"\Delta v(k) = K_p \left( e(k) - e(k-1) \right) + \frac{K_p\,T_s}{T_i}\,e(k)")

add_paragraph(
    "Se ha optado por esta formulación en lugar de la posicional por dos motivos "
    "prácticos. El primero es que ante cambios bruscos de setpoint la forma incremental "
    "no provoca saltos abruptos en la salida, problema conocido en la literatura "
    "industrial como bumpless transfer. El segundo es que simplifica drásticamente la "
    "incorporación del anti-windup, puesto que no requiere mantener una variable de "
    "integración explícita: el efecto integral está distribuido en la acumulación de "
    "los Δv aplicados a lo largo del tiempo.",
)

add_paragraph(
    "Las salidas incrementales Δv de los dos PI pasan por el desacoplador para producir "
    "los incrementos sobre las bombas, que se acumulan y se saturan:",
)

add_formula_box("Saturacion",
    r"\mathbf{u}(k) = \text{sat}\left( \mathbf{u}(k-1) + \mathbf{D}\,\Delta\mathbf{v}(k),\ "
    r"\mathbf{u}_{min},\ \mathbf{u}_{max} \right)")

add_paragraph(
    "La razón para incorporar una estrategia anti-windup es directa: las bombas tienen "
    "límites físicos u_min = 0 y u_max = 2·u_s, lo que implica que el actuador puede "
    "saturarse cuando el error es grande. Sin anti-windup, mientras el actuador permanece "
    "saturado el término integral del PI sigue acumulándose sin efecto real sobre la "
    "planta, puesto que la señal aplicada está fijada por el límite físico. Al desaparecer "
    "la causa de la saturación, el controlador necesita consumir toda la integral "
    "acumulada antes de revertir su acción, lo cual provoca sobrepicos importantes y, en "
    "el peor caso, oscilaciones sostenidas, fenómeno conocido como integrator windup. "
    "En la formulación incremental adoptada, esta estrategia se implementa de forma "
    "natural saturando directamente la señal u(k) antes de aplicarla al actuador: el "
    "incremento que excede el rango simplemente no se acumula porque no se aplica. Esta "
    "variante se conoce como saturación condicional y es la solución más robusta, además "
    "de no requerir parámetros adicionales para su ajuste.",
)

# 4.2.5
add_seccion_nivel3("4.2.5", "Verificación adicional en Simulink")

add_paragraph(
    "Como verificación adicional al análisis basado en scripts de MATLAB, el controlador "
    "PI con desacoplador se replica en el entorno MATLAB/Simulink. La estructura emplea "
    "los bloques de PI discretos del propio entorno, sintonizados con los parámetros "
    "obtenidos por IMC, un bloque que implementa el desacoplador estático y una "
    "representación de la planta no lineal con realimentación de estado. El solver se "
    "configura como ode45 de paso variable con el fin de preservar la precisión numérica "
    "en regiones alejadas del punto de operación, donde las no linealidades son más "
    "pronunciadas. Esta réplica reproduce los resultados de los scripts, lo cual confirma "
    "que las estrategias propuestas son trasladables a entornos de simulación gráfica "
    "ampliamente utilizados en la industria y constituye un primer paso hacia una "
    "eventual implementación en hardware.",
)

# 4.3 Escenario integrado
add_seccion_nivel2("4.3", "Escenario integrado de simulación")

add_paragraph(
    "A diferencia de los trabajos que evalúan el desempeño en escenarios separados (caso "
    "nominal, perturbaciones e incertidumbre), en el presente trabajo se ha diseñado un "
    "escenario integrado que combina, en una sola simulación, los desafíos más "
    "representativos del control multivariable. Esta integración permite comparar el "
    "comportamiento global de ambos controladores en una secuencia operativa coherente "
    "y observar cómo cada uno responde a la sucesión de eventos típicos en una planta "
    "industrial real, donde los eventos no se presentan de manera aislada sino "
    "superpuestos.",
)

add_seccion_nivel3("4.3.1", "Configuración común")

add_paragraph(
    "La planta corresponde al modelo no lineal de los cuatro tanques acoplados desarrollado "
    "en el Capítulo 2, integrado numéricamente mediante ode45. La simulación arranca con "
    "el estado inicial h(0) = [0, 0, 0, 0]ᵀ, es decir, con los tanques vacíos, y con las "
    "bombas apagadas u(0) = [0, 0]ᵀ. Para ambos controladores se adopta el mismo tiempo "
    "de muestreo T_s = 1 s, lo que garantiza que el contraste se realice bajo idénticas "
    "condiciones de discretización. La duración total de la simulación es T_sim = 2000 s "
    "y el punto de operación nominal se mantiene en h₃⁰ = h₄⁰ = 25 cm, mientras que las "
    "restricciones físicas sobre las bombas se fijan en u_min = 0 y u_max = 2·u_s⁰ por "
    "canal.",
)

add_seccion_nivel3("4.3.2", "Trayectoria de referencias y eventos")

add_paragraph(
    "La secuencia de eventos activa, en orden, los aspectos relevantes del control "
    "multivariable. Durante el primer tramo (0 ≤ t < 400 s) el sistema se llena desde "
    "tanques vacíos y se aproxima al punto estacionario, etapa que evalúa la capacidad "
    "del controlador para gestionar el arranque del proceso. En t = 400 s el setpoint "
    "de h₃ cambia de 25 a 30 cm, lo que provoca el primer efecto de acoplamiento cruzado "
    "sobre h₄. En t = 800 s el setpoint de h₄ cambia de 25 a 20 cm, lo que evalúa el "
    "acoplamiento en sentido contrario. En t = 1100 s se activa el ruido gaussiano sobre "
    "las mediciones para evaluar la robustez de cada controlador frente a la presencia "
    "de ruido de los transmisores. Finalmente, en t = 1200 s se produce el evento crítico "
    "del escenario: el setpoint de h₃ regresa a 25 cm mientras el setpoint de h₄ sube a "
    "35 cm, condición que representa un alejamiento del 40% respecto al punto nominal y "
    "combina simultáneamente operación extrema, ruido activo y cambio de referencia "
    "opuesto en el otro lazo.",
)

add_seccion_nivel3("4.3.3", "Inyección de ruido en sensores")

add_paragraph(
    "Para reproducir las condiciones realistas de operación industrial, a partir del "
    "instante t = 1100 s se añade ruido gaussiano blanco a las mediciones de h₃ y h₄ "
    "que ingresan a los controladores:",
)

add_formula_box("Ruido gaussiano",
    r"y_{med,i}(k) = h_i(k) + n_i(k),\quad n_i(k) \sim \mathcal{N}(0, \sigma^2)")

add_paragraph(
    "La desviación estándar se fija en σ = 0.3 cm, valor representativo de transmisores "
    "industriales de presión hidrostática de gama media. La activación del ruido en "
    "t = 1100 s se ubica deliberadamente justo antes del cambio de referencias en t = 1200 s, "
    "lo que permite observar cómo cada controlador responde al ruido cuando además debe "
    "enfrentar el evento crítico del escenario.",
)

add_marcador_imagen("IMAGEN 4.2 — Trayectoria de referencias r_h₃(t) y r_h₄(t). Generada con trayectoria_referencias.m")

# 4.4 Analisis cuantitativo
add_seccion_nivel2("4.4", "Análisis comparativo cuantitativo")

add_seccion_nivel3("4.4.1", "Métricas y ventanas de evaluación")

add_paragraph(
    "Para cada controlador se calculan los seis criterios de desempeño definidos en la "
    "sección 3.2, esto es: IAE, ISE, ITAE, tiempo de establecimiento, sobrepico y esfuerzo "
    "de control. Estos criterios se reportan en dos ventanas temporales diferenciadas. "
    "La primera, denominada métricas globales, se evalúa sobre toda la simulación e "
    "incluye el arranque desde tanques vacíos. La segunda, denominada métricas en "
    "operación normal, se evalúa únicamente a partir de t = 400 s, es decir, una vez "
    "alcanzado el punto estacionario.",
)

add_paragraph(
    "Esta distinción resulta metodológicamente importante puesto que el arranque desde "
    "tanques vacíos representa una fase transitoria de llenado del sistema, no un "
    "escenario de operación normal. Una métrica global puede verse dominada por la "
    "magnitud del error durante el arranque, lo cual oculta el desempeño real del "
    "controlador en régimen operativo, justamente donde se aprecian los efectos del "
    "acoplamiento y la robustez ante perturbaciones. Reportar ambas ventanas proporciona "
    "una visión completa del comportamiento de cada controlador en las distintas fases "
    "del escenario.",
)

add_marcador_imagen("TABLA 4.3 — Métricas globales por controlador. Generada con comparacion_GPC_vs_PID.m")

add_marcador_imagen("TABLA 4.4 — Métricas en operación normal (t ≥ 400 s). Generada con comparacion_GPC_vs_PID.m")

add_seccion_nivel3("4.4.2", "Métrica de acoplamiento cruzado")

add_paragraph(
    "Las métricas integrales clásicas no capturan adecuadamente la magnitud de la "
    "interacción cruzada entre lazos, que es precisamente la característica que se "
    "desea evidenciar al comparar un controlador multivariable con uno descentralizado. "
    "Por esta razón, en el presente trabajo se introduce una métrica específica que "
    "cuantifica cuánto se desvía una salida cuando se modifica únicamente la referencia "
    "de la otra. Para el cambio de referencia en h₃ en t = 400 s, la métrica se define "
    "como la integral del error de h₄ durante una ventana posterior al cambio:",
)

add_formula_box("INT h4",
    r"\text{INT}_{h_4} = \frac{1}{|\Delta r_3|} \int_{t_c}^{t_c + \Delta T} "
    r"\left| e_{h_4}(t) \right|\, dt")

add_paragraph(
    "donde Δr₃ = 5 cm es el cambio nominal del setpoint que provoca el efecto cruzado, "
    "t_c = 400 s el instante del cambio y ΔT = 300 s la ventana de observación. De manera "
    "análoga se define INT_{h_3} para el cambio de referencia en h₄ en t = 800 s. Valores "
    "bajos de INT indican un desacoplamiento efectivo entre los lazos, mientras que "
    "valores altos evidencian una interacción cruzada significativa.",
)

add_marcador_imagen("TABLA 4.5 — Métrica de acoplamiento cruzado. Generada con comparacion_GPC_vs_PID.m")

add_seccion_nivel3("4.4.3", "Resultados gráficos")

add_marcador_imagen("IMAGEN 4.3 — Respuesta comparativa de h₃ durante toda la simulación. Generada con comparacion_GPC_vs_PID.m")
add_marcador_imagen("IMAGEN 4.4 — Respuesta comparativa de h₄ durante toda la simulación. Generada con comparacion_GPC_vs_PID.m")
add_marcador_imagen("IMAGEN 4.5 — Señales de control u₁ y u₂. Generada con comparacion_GPC_vs_PID.m")

# 4.5 Discusion
add_seccion_nivel2("4.5", "Discusión y validación de la hipótesis")

add_seccion_nivel3("4.5.1", "Comportamiento ante acoplamiento cruzado")

add_paragraph(
    "El PI descentralizado, aun con desacoplador estático, presenta perturbaciones notorias "
    "en una salida cuando cambia el setpoint de la otra. Esto se debe a que el desacoplador "
    "estático cancela la interacción únicamente en estado estacionario, mientras que "
    "durante el régimen transitorio los acoples cruzados dinámicos no son compensados. El "
    "GPC, en cambio, al considerar la dinámica completa del sistema en su predicción a N "
    "pasos, anticipa el efecto del acople y coordina simultáneamente las dos entradas para "
    "minimizarlo desde el primer paso. La métrica INT reportada en la Tabla 4.5 cuantifica "
    "esta ventaja del GPC frente al esquema clásico.",
)

add_seccion_nivel3("4.5.2", "Robustez ante ruido de medición")

add_paragraph(
    "Tras la inyección de ruido en t = 1100 s, ambos controladores transmiten parte del "
    "ruido a las señales de control. Sin embargo, el GPC exhibe una menor amplificación "
    "del ruido en las bombas que el PI con desacoplador, debido a la ponderación λ del "
    "esfuerzo de control incluida en su función de costo, que penaliza explícitamente las "
    "variaciones bruscas de la señal manipulada. Esto se refleja en una menor variación "
    "de u durante el último tramo de la simulación, característica relevante para "
    "preservar la vida útil de los actuadores en una aplicación industrial real.",
)

add_seccion_nivel3("4.5.3", "Operación lejos del punto de linealización")

add_paragraph(
    "El tramo crítico del escenario corresponde a t ≥ 1200 s, cuando el setpoint de h₄ "
    "se establece en 35 cm. En esta región, la planta no lineal presenta dinámicas "
    "significativamente distintas a las consideradas en la sintonización IMC del PI, dado "
    "que las constantes de tiempo de los tanques dependen de √h y, por tanto, varían con "
    "el nivel de operación. Como consecuencia, el PI descentralizado no alcanza el "
    "setpoint de 35 cm y exhibe un error en estado estacionario persistente. La "
    "sintonización IMC con λ_imc = τ/3, calculada para h = 25 cm, no proporciona la "
    "ganancia adecuada para esta región operativa, mientras que el desacoplador estático "
    "tampoco compensa adecuadamente puesto que sus coeficientes asumen la matriz de "
    "ganancias DC del punto nominal. El GPC, en cambio, sí alcanza el setpoint de 35 cm, "
    "aunque con un transitorio más lento que en la región nominal. Esto se debe a que su "
    "capacidad de predicción permite anticipar el efecto de la entrada acumulada sobre el "
    "horizonte futuro, compensando parcialmente las no linealidades del modelo lineal "
    "interno.",
)

add_paragraph(
    "Este resultado constituye una validación experimental contundente de la hipótesis "
    "principal del trabajo: el control predictivo extiende la región de operación "
    "admisible del sistema más allá del entorno inmediato del punto de linealización, "
    "mientras que el PI con desacoplador queda restringido a una vecindad estrecha de su "
    "sintonización original.",
)

add_seccion_nivel3("4.5.4", "Trade-offs identificados")

add_paragraph(
    "La comparación no es absoluta y conviene reconocer que las ventajas del GPC vienen "
    "acompañadas de costos asociados. El primero es un costo computacional notablemente "
    "superior, puesto que la resolución del problema de optimización cuadrática en cada "
    "periodo de muestreo es considerablemente más exigente que la evaluación recursiva de "
    "un PI. Para tiempos de muestreo del orden del segundo, como el adoptado en el "
    "presente trabajo, esto no representa una limitación en hardware industrial moderno; "
    "en cambio, en procesos rápidos del orden de milisegundos sí debería evaluarse caso "
    "por caso. El segundo es una mayor complejidad de implementación, puesto que el GPC "
    "requiere infraestructura matemática (modelo, optimizador) ausente en el PI, lo que "
    "implica una curva de aprendizaje para el personal técnico y mayor dependencia de "
    "software especializado. El tercero es una dependencia más estricta del modelo del "
    "proceso, pues tanto el GPC como el desacoplador requieren conocer el modelo del "
    "proceso, pero el GPC lo emplea de manera más robusta (proyección al futuro) que el "
    "desacoplador (inversión algebraica del punto nominal). La cuantificación de estos "
    "costos permite responder de manera fundamentada a la pregunta de cuándo se justifica "
    "implementar un GPC en lugar de un PI con desacoplador en una aplicación industrial.",
)

# 4.6 Conclusiones del capitulo
add_seccion_nivel2("4.6", "Conclusiones del capítulo")

add_paragraph(
    "En el presente capítulo se ha desarrollado un análisis comparativo del controlador "
    "predictivo generalizado (GPC) diseñado en el Capítulo 3 frente a un controlador PI "
    "descentralizado con desacoplador estático sintonizado por Internal Model Control, "
    "aplicado al sistema hidráulico de cuatro tanques acoplados. La selección del "
    "controlador de contraste responde a la práctica industrial estándar: se ha optado "
    "por un PI en lugar de un PID debido a la sensibilidad del término derivativo al "
    "ruido en dinámicas dominantemente de primer orden, y se ha incorporado un desacoplador "
    "estático con el fin de reflejar la mejor configuración clásica practicable en un "
    "sistema TITO acoplado.",
)

add_paragraph(
    "La comparación se sustenta en un escenario integrado que combina, en una sola "
    "simulación, los principales desafíos del control multivariable: arranque del sistema "
    "desde tanques vacíos, cambios secuenciales y simultáneos de referencia, inyección de "
    "ruido gaussiano en los sensores y operación significativamente alejada del punto de "
    "linealización. Los resultados cuantitativos, respaldados por los seis criterios de "
    "desempeño del Capítulo 3 y la métrica específica de acoplamiento cruzado introducida "
    "en este trabajo, evidencian que el GPC reduce significativamente la interacción "
    "cruzada entre lazos, presenta menor amplificación del ruido en las señales de control "
    "y extiende la región de operación admisible más allá del entorno inmediato del punto "
    "de linealización.",
)

add_paragraph(
    "El caso más representativo se observa en el tramo final del escenario (t ≥ 1200 s), "
    "donde el setpoint de h₄ se establece en 35 cm: el PI con desacoplador no alcanza esta "
    "referencia debido a las no linealidades del modelo y a la sintonización fijada para "
    "el punto nominal, mientras que el GPC sí lo logra gracias a su capacidad predictiva. "
    "Este resultado valida empíricamente la hipótesis principal del trabajo y establece, "
    "junto con los trade-offs identificados (mayor costo computacional, mayor complejidad "
    "de implementación y mayor dependencia del modelo), los criterios bajo los cuales el "
    "GPC justifica su adopción frente al PI con desacoplador en aplicaciones industriales "
    "de control multivariable.",
)

# ============================================================
# CAPITULO 5: CONCLUSIONES
# ============================================================
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run("Conclusiones")
r.font.name = FONT; r.font.size = Pt(14); r.bold = True

add_paragraph(
    "El presente trabajo desarrolló el diseño, la implementación en simulación y el "
    "análisis comparativo de un controlador Predictivo Generalizado (GPC) multivariable "
    "aplicado al sistema hidráulico de cuatro tanques acoplados propuesto por Johansson, "
    "empleando como referencia comparativa un controlador PI descentralizado con "
    "desacoplador estático sintonizado por Internal Model Control. A partir de los "
    "resultados obtenidos a lo largo de los cuatro capítulos, se establecen las "
    "siguientes conclusiones:",
)

conclusiones = [
    "Se caracterizó exitosamente la dinámica multivariable del sistema de cuatro tanques "
    "acoplados, identificando los caminos directos y cruzados entre entradas y salidas, "
    "y verificando el régimen de fase mínima correspondiente al banco de pruebas del "
    "Laboratorio de Control Avanzado de la PUCP mediante la condición γ₁ + γ₂ > 1.",

    "Se obtuvo el modelo linealizado del sistema en espacio de estados y su matriz de "
    "funciones de transferencia MIMO en torno al punto de operación nominal "
    "h₃⁰ = h₄⁰ = 25 cm. La validación mediante la métrica FIT% confirmó una alta fidelidad "
    "del modelo lineal dentro de una vecindad de ±10% del punto de operación, y evidenció "
    "su degradación progresiva a medida que el sistema se aleja de dicha vecindad.",

    "Se diseñó el controlador GPC MIMO bajo la formulación CARIMA con ecuaciones "
    "diofánticas, adoptando un tiempo de muestreo T_s = 1 s, un horizonte de predicción "
    "N = 50 y un horizonte de control N_u = 5. Los pesos de la función de costo "
    "(δ = [10, 10] y λ ≈ 0.0077) se obtuvieron mediante un análisis comparativo de cuatro "
    "métodos de sintonización (Clarke-Mohtadi, Shridhar-Cooper, PSO y Nelder-Mead) "
    "evaluados con un score combinado de seis métricas.",

    "El tratamiento de las restricciones físicas sobre las bombas se incorporó de manera "
    "explícita mediante una formulación de programación cuadrática (QP), resuelta en cada "
    "periodo de muestreo con el algoritmo active-set. Esta característica constituye una "
    "ventaja estructural del GPC frente a las estrategias clásicas, que sólo pueden manejar "
    "las restricciones mediante saturación posterior.",

    "El análisis comparativo del Capítulo 4 evidenció que el GPC reduce significativamente "
    "la interacción cruzada entre lazos respecto al PI con desacoplador. La métrica "
    "específica de acoplamiento INT introducida en este trabajo cuantifica esta ventaja "
    "y establece un criterio objetivo para el contraste entre estrategias multivariables "
    "y descentralizadas.",

    "El resultado más contundente se observó en el tramo t ≥ 1200 s del escenario "
    "integrado, cuando el setpoint de h₄ se establece en 35 cm (40% por encima del punto "
    "nominal). En estas condiciones, el PI con desacoplador no alcanza el setpoint debido "
    "a la variación de las constantes de tiempo con √h y a la pérdida de efectividad del "
    "desacoplador fuera del punto de operación nominal, mientras que el GPC sí lo alcanza "
    "gracias a su capacidad predictiva. Este resultado valida empíricamente la hipótesis "
    "principal del trabajo.",

    "Los trade-offs identificados —mayor costo computacional, mayor complejidad de "
    "implementación y mayor dependencia del modelo— no comprometen la aplicabilidad del "
    "GPC en el escenario estudiado, dado que el tiempo de muestreo adoptado (T_s = 1 s) "
    "es holgado para hardware industrial moderno. En procesos con dinámicas más rápidas, "
    "estos costos deberían evaluarse caso por caso.",

    "La implementación en Simulink de ambos controladores reprodujo los resultados de "
    "los scripts de MATLAB, demostrando la reproducibilidad de las estrategias propuestas "
    "en entornos de simulación gráfica ampliamente utilizados en la industria y sentando "
    "las bases para una eventual validación experimental sobre la planta piloto del "
    "Laboratorio de Control Avanzado de la PUCP.",
]

for i, c in enumerate(conclusiones):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i+1}. {c}")
    r.font.name = FONT; r.font.size = Pt(12)

# ============================================================
# RECOMENDACIONES (Trabajo futuro)
# ============================================================
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(24)
r = p.add_run("Recomendaciones")
r.font.name = FONT; r.font.size = Pt(14); r.bold = True

add_paragraph(
    "A partir del trabajo desarrollado y de las limitaciones identificadas, se plantean "
    "las siguientes recomendaciones para líneas futuras de investigación:",
)

recomendaciones = [
    "Validar experimentalmente el controlador GPC diseñado sobre la planta piloto del "
    "Laboratorio de Control Avanzado de la PUCP, incorporando las particularidades del "
    "hardware real: retardos de comunicación, resolución de los actuadores y "
    "características específicas de los transmisores de nivel.",

    "Estudiar variantes del GPC que permitan manejar de manera explícita las no "
    "linealidades del proceso, como el Nonlinear Model Predictive Control (NMPC) o el "
    "GPC adaptativo con actualización del modelo en línea. Estas variantes podrían "
    "extender aún más la región de operación admisible del sistema.",

    "Evaluar la robustez del controlador frente a incertidumbre paramétrica y "
    "envejecimiento del proceso mediante análisis de sensibilidad y experimentos "
    "Monte Carlo sobre los parámetros físicos del modelo (áreas, coeficientes de "
    "descarga, ganancias de las bombas).",

    "Integrar el controlador desarrollado con plataformas SCADA o MES industriales, "
    "de modo que sea posible evaluar su desempeño en un entorno operativo cercano al "
    "de la industria real, incluyendo la gestión de alarmas y la comunicación con "
    "sistemas de supervisión.",

    "Extender el análisis comparativo a otros esquemas de control avanzado, como el "
    "control robusto H∞ o el control por modos deslizantes, con el fin de completar el "
    "mapa de estrategias aplicables al sistema de cuatro tanques y establecer un "
    "referente amplio para la selección de la estrategia más adecuada en aplicaciones "
    "industriales similares.",
]

for i, r_txt in enumerate(recomendaciones):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = INDENT
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{i+1}. {r_txt}")
    r.font.name = FONT; r.font.size = Pt(12)

# ============================================================
# BIBLIOGRAFIA (formato APA)
# ============================================================
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run("Bibliografía")
r.font.name = FONT; r.font.size = Pt(14); r.bold = True

bibs = [
    "[1] Laubwald, E. (2005). Coupled tank system. Control Systems Principles, 1–8.",

    "[2] Yuan, W. (2023). Mathematical Model Analysis and Control Strategy. 2023 International "
    "Conference on Mechatronics, Control and Robotics (ICMCR), 94–97. Jeju, Korea: IEEE. "
    "https://doi.org/10.1109/ICMCR56776.2023.10181040",

    "[3] Short, M., & Selvakumar, A. (2020). Non-Linear Tank Level Control for Industrial "
    "Applications. Applied Mathematics, 11, 876–889.",

    "[4] Numsomran, V., Tipsuwanporn, V., & Tirasesth, K. (2008). Modeling of the Modified "
    "Quadruple-Tank Process. 2008 SICE Annual Conference, 818–823. Chofu, Japan: IEEE. "
    "https://doi.org/10.1109/SICE.2008.4654768",

    "[5] Azam, S. N. M., & Jørgensen, J. B. (2015). Modeling and simulation of a modified "
    "quadruple tank system. 2015 IEEE International Conference on Control System, Computing "
    "and Engineering (ICCSCE), 365–370. Penang, Malaysia: IEEE. "
    "https://doi.org/10.1109/ICCSCE.2015.7482213",

    "[6] Yu, Y., Yang, H., Wan, S., Liu, Q., & Yan, J. (2024). Un método de control cooperativo "
    "y su aplicación para sistemas acoplados multivariables en serie. Scientific Reports, 14. "
    "https://doi.org/10.1038/s41598-024-63169-7",

    "[7] Albertos, P., & Sala, A. (2004). Multivariable Control Systems: An Engineering "
    "Approach. London, UK: Springer. https://doi.org/10.1007/b97506",

    "[8] Johansson, K. H. (2000). The quadruple-tank process: a multivariable laboratory "
    "process with an adjustable zero. IEEE Transactions on Control Systems Technology, 8(3), "
    "456–465. https://doi.org/10.1109/87.845876",

    "[9] Pugliese, L., De Oliveira, T., Da Silva, D., Rodor, F., Braga, R., & Amorim, G. "
    "(2022). Modelado y desarrollo de una planta didáctica de bajo coste para la enseñanza en "
    "sistemas multivariables. Research, Society and Development, 11(7). "
    "https://doi.org/10.33448/rsd-v11i7.30249",

    "[10] Sánchez Zurita, V. A. (2018). Diseño de un sistema de control predictivo multivariable "
    "aplicado a un proceso de cuatro tanques acoplados. Tesis de Maestría, Pontificia "
    "Universidad Católica del Perú, Lima, Perú.",

    "[11] Gouta, H., Haysam Al-Ashek, W., & Saad, B. (2022). Anti-disturbance composite "
    "tracking control for a coupled two-tank MIMO process with experimental studies. Automatika, "
    "63(3), 593–604. https://doi.org/10.1080/00051144.2022.2059207",

    "[12] Tang, J., Zhao, S., Fu, Q., Liu, Z., & He, W. (2021). Adaptive fault-tolerant control "
    "for a three-tank system with height and rate constraints. 2021 China Automation Congress "
    "(CAC), 4020–4024. Beijing, China: IEEE. https://doi.org/10.1109/CAC53003.2021.9728310",

    "[13] Abushokor, A., & Amrr, S. M. (2025). Model-Free Adaptive Time-Delay-Based Estimation "
    "Control for Input-Saturated Coupled Tank System: Experimental Validation. IEEE Transactions "
    "on Automation Science and Engineering, 22, 19340–19351. "
    "https://doi.org/10.1109/TASE.2025.3594752",

    "[14] Santana, H. G., Coelho, S. de S., & de Almeida, O. de M. (2018). Application of "
    "Multivariable PID Controllers in a Coupled Tank System. 2018 13th IEEE International "
    "Conference on Industry Applications (INDUSCON), 664–671. Sao Paulo, Brazil: IEEE. "
    "https://doi.org/10.1109/INDUSCON.2018.8627072",

    "[15] Cartes, D., & Wu, L. (2005). Experimental evaluation of adaptive three-tank level "
    "control. ISA Transactions, 44(2), 283–293. https://doi.org/10.1016/S0019-0578(07)60181-5",

    "[16] Choudhary, P. K., Raj, P., & Das, D. K. (2024). Controller Design for Decoupled "
    "Two-Input Two-Output Coupled Tank System. 2024 IEEE International Conference on Smart "
    "Power Control and Renewable Energy (ICSPCRE), 1–6. Rourkela, India: IEEE. "
    "https://doi.org/10.1109/ICSPCRE62303.2024.10675217",

    "[17] Rivera, D. E., Morari, M., & Skogestad, S. (1986). Internal Model Control: PID "
    "Controller Design. Industrial & Engineering Chemistry Process Design and Development, "
    "25(1), 252–265. https://doi.org/10.1021/i200032a041",

    "[18] Åström, K. J., & Hägglund, T. (2006). Advanced PID Control. Research Triangle Park, "
    "NC: ISA - The Instrumentation, Systems, and Automation Society.",

    "[19] Skogestad, S., & Postlethwaite, I. (2005). Multivariable Feedback Control: Analysis "
    "and Design (2nd ed.). Chichester, UK: John Wiley & Sons.",

    "[20] Shridhar, R., & Cooper, D. J. (1997). A Tuning Strategy for Unconstrained Multivariable "
    "Model Predictive Control. Industrial & Engineering Chemistry Research, 37(10), 4003–4016. "
    "https://doi.org/10.1021/ie980202s",

    "[21] Eberhart, R., & Kennedy, J. (1995). A new optimizer using particle swarm theory. "
    "MHS'95 Proceedings of the Sixth International Symposium on Micro Machine and Human "
    "Science, 39–43. Nagoya, Japan: IEEE. https://doi.org/10.1109/MHS.1995.494215",

    "[22] Åström, K. J., & Wittenmark, B. (1997). Computer-Controlled Systems: Theory and "
    "Design (3rd ed.). Upper Saddle River, NJ: Prentice Hall.",

    "[23] Camacho, E. F., & Bordons, C. (2007). Model Predictive Control (2nd ed.). London, "
    "UK: Springer. https://doi.org/10.1007/978-0-85729-398-5",

    "[24] Clarke, D. W. (1988). Application of Generalized Predictive Control to Industrial "
    "Processes. IEEE Control Systems Magazine, 8(2), 49–55. "
    "https://doi.org/10.1109/37.1961",

    "[25] Cheng, Y. (2007). Predicción de j pasos adelante basada en el modelo CARMA para "
    "sistemas MIMO. Frontiers of Electrical and Electronic Engineering in China, 2, 99–103. "
    "https://doi.org/10.1007/s11460-007-0018-7",

    "[26] Nelder, J. A., & Mead, R. (1965). A Simplex Method for Function Minimization. The "
    "Computer Journal, 7(4), 308–313. https://doi.org/10.1093/comjnl/7.4.308",

    "[27] Ogata, K. (2010). Modern Control Engineering (5th ed.). Upper Saddle River, NJ: "
    "Prentice Hall.",
]

for b in bibs:
    add_bibliografia_entry(b)

# ============================================================
# ANEXOS CON CODIGOS MATLAB
# ============================================================

anexos = [
    ("A", "Código MATLAB del controlador GPC",
     "El presente anexo contiene el código MATLAB del controlador Predictivo Generalizado "
     "(GPC) diseñado en el Capítulo 3, aplicado sobre la planta no lineal de cuatro tanques "
     "acoplados. El script implementa la formulación CARIMA + ecuaciones diofánticas, "
     "la construcción de la matriz dinámica G por bloques, la matriz de ponderación de la "
     "función de costo, el tratamiento de restricciones mediante programación cuadrática "
     "y el bucle de control con horizonte deslizante.",
     [("controlador_GPC.m", "controlador_GPC.m")]),

    ("B", "Código MATLAB del controlador PI con desacoplador estático",
     "El presente anexo contiene el código MATLAB del controlador PI descentralizado con "
     "desacoplador estático desarrollado en el Capítulo 4. El script implementa la "
     "sintonización IMC de los dos PI, el cálculo de los coeficientes del desacoplador a "
     "partir de la matriz de ganancias DC del sistema, la forma incremental con anti-windup "
     "por saturación condicional y el bucle de control sobre la planta no lineal.",
     [("controlador_PID.m", "controlador_PID.m")]),

    ("C", "Código MATLAB de la comparación GPC vs PI + Desacoplador",
     "El presente anexo contiene el código MATLAB que ejecuta el análisis comparativo "
     "central del Capítulo 4. El script simula ambos controladores sobre el mismo escenario "
     "integrado, calcula las métricas globales, las métricas en operación normal, la métrica "
     "de acoplamiento cruzado INT y genera las figuras comparativas.",
     [("comparacion_GPC_vs_PID.m", "comparacion_GPC_vs_PID.m")]),

    ("D", "Códigos MATLAB de los análisis auxiliares del Capítulo 3 y 4",
     "El presente anexo agrupa los códigos MATLAB de los análisis auxiliares. El primero "
     "corresponde al análisis comparativo de los cuatro métodos de sintonización del GPC "
     "descritos en el Capítulo 3. El segundo justifica la elección de λ_imc = τ/3 para la "
     "sintonización del PI. El tercero genera la trayectoria de referencias del escenario "
     "integrado del Capítulo 4.",
     [("analisis_sintonizacion_GPC.m", "analisis_sintonizacion_GPC.m"),
      ("analisis_lambda_imc.m",       "analisis_lambda_imc.m"),
      ("trayectoria_referencias.m",   "trayectoria_referencias.m")]),

    ("E", "Códigos MATLAB Function de los bloques de Simulink",
     "El presente anexo contiene los códigos de los bloques MATLAB Function empleados en "
     "los modelos Simulink de los controladores GPC y PI con desacoplador. Corresponden "
     "respectivamente a: el paso de control del GPC (formulación QP con coder.extrinsic "
     "para quadprog), el desacoplador estático de Skogestad y la planta no lineal integrada "
     "mediante Euler subdividido.",
     [("gpc_step_simulink.m",         "gpc_step_simulink.m"),
      ("desacoplador_simulink.m",     "desacoplador_simulink.m"),
      ("planta_no_lineal_simulink.m", "planta_no_lineal_simulink.m")]),
]

for letra, titulo, intro, archivos in anexos:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(f"Anexo {letra}")
    r.font.name = FONT; r.font.size = Pt(14); r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run(titulo)
    r.font.name = FONT; r.font.size = Pt(14); r.bold = True

    add_paragraph(intro)
    doc.add_paragraph()

    for nombre, ruta_rel in archivos:
        add_codigo_matlab(nombre, os.path.join(BASE, ruta_rel))
        doc.add_paragraph()

# ============================================================
# GUARDAR
# ============================================================
OUT = os.path.join(BASE, "TESIS_20212444_secciones_nuevas.docx")
doc.save(OUT)
print(f"Documento generado: {OUT}")
print(f"Slides/parrafos totales: {len(doc.paragraphs)}")
