"""Verifica que el docx final tenga las secciones en orden correcto."""
from docx import Document
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc = Document(r"C:\Users\HP\OneDrive\Escritorio\edwin\TESIS_20212444_FINAL.docx")

marcadores = [
    'INTRODUCCIÓN', 'Introducción',
    '1. Estudio', '1.5', '1.5.1', '1.5.2',
    '2. Fundamentos',
    '3. Diseño',
    'Capítulo 4', 'CAPITULO 4', 'CAPÍTULO 4',
    '4.1', '4.2', '4.3', '4.4', '4.5', '4.6',
    'Conclusiones', 'Recomendaciones', 'Bibliografía',
    'Anexo A', 'Anexo B', 'Anexo C', 'Anexo D', 'Anexo E',
    '[1]', '[10]', '[20]', '[27]',
]

encontrados = {}
for i, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    for m in marcadores:
        if txt.startswith(m):
            if m not in encontrados:
                encontrados[m] = []
            encontrados[m].append((i, txt[:80]))

print(f"Total parrafos: {len(doc.paragraphs)}\n")
for m in marcadores:
    if m in encontrados:
        print(f"[OK] {m:20s} -> {len(encontrados[m])} ocurrencias, primera en [{encontrados[m][0][0]:04d}]")
    else:
        print(f"[XX] {m:20s} -> NO ENCONTRADO")

# Mostrar los primeros parrafos despues del titulo INTRODUCCION
print("\n=== Contenido despues de INTRODUCCION ===")
for m in ['INTRODUCCIÓN']:
    if m in encontrados:
        idx = encontrados[m][0][0]
        for i in range(idx, min(idx+6, len(doc.paragraphs))):
            print(f"[{i:04d}] {doc.paragraphs[i].text.strip()[:150]}")

print("\n=== Contenido despues de Capitulo 4 ===")
for m in ['Capítulo 4', 'CAPITULO 4', 'CAPÍTULO 4']:
    if m in encontrados:
        idx = encontrados[m][0][0]
        for i in range(idx, min(idx+8, len(doc.paragraphs))):
            print(f"[{i:04d}] {doc.paragraphs[i].text.strip()[:150]}")
        break

print("\n=== Contenido despues de Conclusiones ===")
for m in ['Conclusiones']:
    if m in encontrados:
        for idx, txt in encontrados[m]:
            print(f"\n>> Ocurrencia en [{idx}]:")
            for i in range(idx, min(idx+4, len(doc.paragraphs))):
                print(f"  [{i:04d}] {doc.paragraphs[i].text.strip()[:150]}")
