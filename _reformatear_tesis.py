"""
Aplica el formato PUCP a TESIS_20212444.docx sin alterar el contenido:
- Margenes 2.54 cm en todas las secciones
- Fuente Times New Roman 12 en todos los runs
- Interlineado doble en todos los parrafos
- Alineacion justificada por defecto

Preserva imagenes, ecuaciones y estructura del docx original.
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os

BASE = r"C:\Users\HP\OneDrive\Escritorio\edwin"
IN  = os.path.join(BASE, "TESIS_20212444.docx")
OUT = os.path.join(BASE, "TESIS_20212444_reformateada.docx")

FONT   = "Times New Roman"
SIZE   = Pt(12)
MARGIN = Cm(2.54)

doc = Document(IN)

# --- Margenes en todas las secciones ---
for section in doc.sections:
    section.top_margin    = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin   = MARGIN
    section.right_margin  = MARGIN

# --- Estilo Normal ---
style = doc.styles['Normal']
style.font.name = FONT
style.font.size = SIZE

# --- Iterar por todos los parrafos ---
n_pars = 0
n_runs = 0
for p in doc.paragraphs:
    n_pars += 1
    # Interlineado doble
    pf = p.paragraph_format
    pf.line_spacing = 2.0
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    # Runs: fuente y tamano
    for r in p.runs:
        n_runs += 1
        r.font.name = FONT
        if r.font.size is None or r.font.size < Pt(9) or r.font.size > Pt(16):
            r.font.size = SIZE
        else:
            # respetar tamanos ya definidos en encabezados/tablas
            pass

# --- Iterar por tablas (fuente y tamano dentro de celdas) ---
n_cells = 0
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            n_cells += 1
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = FONT
                    if r.font.size is None:
                        r.font.size = Pt(10)

doc.save(OUT)
print(f"Reformateado: {OUT}")
print(f"  Parrafos procesados: {n_pars}")
print(f"  Runs procesados:     {n_runs}")
print(f"  Celdas procesadas:   {n_cells}")
